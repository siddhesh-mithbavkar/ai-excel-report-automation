from docx import Document
import re


def markdown_to_docx_runs(paragraph, text: str):
    """
    Convert **bold** markdown into proper bold runs inside a docx paragraph.
    """
    # Split by **...**
    parts = re.split(r"(\*\*.+?\*\*)", text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            # Bold text inside **
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            paragraph.add_run(part)


def generate_docx(input_file, output_file):
    # Read AI summary
    with open(input_file, "r", encoding="utf-8") as f:
        raw_text = f.read()

    doc = Document()

    # Add title
    doc.add_heading("AI-Generated Sales Report Summary", level=1)

    # Split into blocks by blank lines
    blocks = raw_text.split("\n\n")

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # Create a paragraph
        para = doc.add_paragraph()
        # Support line breaks and bold markdown
        lines = block.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                para.add_run("\n")  # manual line break
            markdown_to_docx_runs(para, line)

    doc.save(output_file)
    print("Word report generated:", output_file)


if __name__ == "__main__":
    generate_docx("ai_summary.txt", "AI_Report.docx")
